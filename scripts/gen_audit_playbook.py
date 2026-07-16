"""Generate audit analysis + playbook .docx"""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

# ── Cover ───────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading("Semantická analýza auditu + korelace se zdrojovým kódem", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
sub = doc.add_paragraph("linkedin-mcp-custom v0.1.0 — Playbook pro debugging, produkční stav & publish-ready")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(14)
sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
doc.add_paragraph()
meta = doc.add_paragraph(f"Datum: {datetime.now().strftime('%d. %m. %Y')}   |   Zdroj: cross-LLM audit od Claude (16.07.2026)   |   Commit: bd5b5dc + večerní P0-P2 featy")
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.runs[0].font.size = Pt(9)
meta.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
doc.add_page_break()

# ── 1. Semantická analýza auditu ──────────────────────────────────────
doc.add_heading("1. Semantická analýza auditu — shrnutí a interpretace", level=1)

doc.add_heading("1.1 Povaha auditu", level=2)
doc.add_paragraph(
    "Audit od Claude je cross-LLM analýza provedená na základě audit_prompt_linkedin_mcp_custom.md "
    "(commit bd5b5dc) a error_analysis_20260715_194245.json. Nejde o audit živého kódu, ale o "
    "nezávislé review dodaných podkladů. Claude neměl přístup k aktuálnímu source code — audit "
    "vychází výhradně z Appendixu A audit promptu (full source snapshot při commitu bd5b5dc) "
    "a z error analysis JSON."
)

doc.add_paragraph(
    "Důsledek: audit nereflektuje večerní změny (P0 sequential fallback, P1 progressive timeout, "
    "P1 HTTP logging, P2 pipeline backoff), které byly implementovány po commitu bd5b5dc. "
    "Některá negativní zjištění (F-01, F-03, chybějící HTTP logging, absence fallbacku) "
    "již byla adresována."
)

doc.add_heading("1.2 Kategorizace nálezů", level=2)
doc.add_paragraph(
    "Audit identifikuje 19 nálezů (F-01 až F-19) ve 4 kategoriích závažnosti: "
    "2 kritické (F-01, F-02), 6 závažných (F-03 až F-08), 7 menších (F-09 až F-15), "
    "2 publish-readiness (F-16, F-17) a 1 info (F-19)."
)

doc.add_heading("1.3 Tři nejdůležitější zjištění auditu", level=2)

findings = [
    ("F-02: Sdílený per-job timeout",
     "Kritické — audit odhaluje, že asyncio.wait_for(timeout=120) není ukotven k okamžiku "
     "získání semaforu, ale k začátku celé per-job fáze. 29 úloh má identickou dobu 120.0s, "
     "což je silný signál jediného sdíleného časovače. Toto zjištění nebylo v původní error "
     "analysis identifikováno a je nejcennějším přínosem auditu."),
    ("F-01: LinkedIn rate-limiting",
     "Potvrzeno auditem — binární charakter selhání (přesně 2 úspěchy, pak 100% timeout) a "
     "identická doba 15.0s u všech 17 selhání odpovídá server-side gate, ne postupné degradaci. "
     "Hypotéza B (vyčerpání Chromia) je auditem oslabena."),
    ("F-04: Verdiktová nesrovnalost",
     "Skóre 35.9 (nad prahem hranicni 30.0) je označeno jako NESLEDOVAT místo HRANICNI. "
     "Audit správně upozorňuje, že jde o jedinou nesrovnalost ze 4 úloh — může jít o edge-case "
     "bug, zaokrouhlení, nebo artefakt sestavení error analysis JSON."),
]

for title_text, body_text in findings:
    p = doc.add_paragraph()
    run = p.add_run(f"{title_text}: ")
    run.bold = True
    p.add_run(body_text)

doc.add_heading("1.4 Co audit neviděl (večerní featy)", level=2)
p = doc.add_paragraph(
    "Audit hodnotí kód při commitu bd5b5dc. Následující změny provedené 15.07. večer nejsou "
    "v auditu reflektovány, a tudíž některá negativní zjištění již neplatí:"
)
items = [
    "P0 Sequential fallback — adresuje F-03 (chybějící sekvenční fallback)",
    "P1 HTTP status logging v _retry_goto — adresuje F-01 (diagnostické logování HTTP statusů)",
    "P1 Progressive timeout 15s→30s — adresuje F-01 (adaptivní timeout)",
    "P2 Pipeline-level backoff (3× timeout → 30s pauza) — adresuje F-01/F-03",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("1.5 Co audit neodhalil (vlastní nálezy z korelace)", level=2)
items2 = [
    "_track_nav_task() je dead code — definováno, ale nikde voláno",
    "create_page() nedědí route interception (resource blocking) — pool pages nemají blokování",
    "Žádný lock okolo singleton browser creation — race condition při souběžných voláních",
    "get_saved_jobs() (MCP tool) nemá timeout — může viset neomezeně dlouho",
    "VERDICT_ICON duplikován v kb_writer.py a report_generator.py — DRY",
    "auth.py race: check_cached_auth a _mark_auth_ok nejsou atomické",
]
for item in items2:
    doc.add_paragraph(item, style="List Bullet")

doc.add_page_break()

# ── 2. Architekturní korelace ──────────────────────────────────────────
doc.add_heading("2. Architekturní korelace — audit vs. source code", level=1)

doc.add_paragraph(
    "Následující tabulka mapuje každý nález auditu (F-01 až F-19) na konkrétní "
    "soubor/y a řádky source code, ověřuje platnost, a hodnotí závažnost."
)

# F-inding table
table = doc.add_table(rows=1, cols=5)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = table.rows[0].cells
hdr[0].text = "ID"
hdr[1].text = "Lokalizace v kódu"
hdr[2].text = "Audit tvrdí"
hdr[3].text = "Korelace (pravda/ nepravda / částečně)"
hdr[4].text = "Status"

rows_data = [
    ("F-01", "extractor.py:79-120 (_retry_goto)\nscraping/utils.py (is_rate_limited)",
     "LinkedIn blokuje /jobs/view/ po ~2 úspěšných",
     "PRAVDA — potvrzeno 3+ běhy; typické pro server-side rate-limit gate",
     "⚠️ Mitigováno (P0-P2)"),
    ("F-02", "run_pipeline.py:410-426 (_run_job_with_timeout)",
     "wait_for(timeout=120) je sdílený časovač, ne per-job",
     "PRAVDA — 29× 120.0s je statisticky nemožné pro nezávislé časovače. wait_for obaluje celý task včetně čekání na semafor",
     "🔴 Nutná oprava"),
    ("F-03", "run_pipeline.py:298-407 (semaphore + _process_one_job)",
     "Chybí pipeline-level backoff",
     "ČÁSTEČNĚ — P2 backoff již implementován (3× timeout → 30s pauza), ale chybí rolling-window",
     "⚠️ Částečně hotovo"),
    ("F-04", "analysis/scorer.py:112-120 (_determine_verdict)\nerror_analysis JSON",
     "35.9 ≥ 30.0 → má být HRANICNI, ne NESLEDOVAT",
     "NEPOTVRZENO — je třeba ověřit v surovém pipeline JSON. Prahy v scorer.py jsou [65,50,40,0], nikoli [70,45,30] z config_snapshot",
     "🔴 Nutná verifikace"),
    ("F-05", "analysis/growth.py:9-20 (growth_score)",
     "growth_score bere jen company name, ignoruje text",
     "PRAVDA — signatura je growth_score(company), ne growth_score(company, text). 20 lines, 2 statické seznamy",
     "⏳ Plánováno (P2-3)"),
    ("F-06", "— (chybějící testy)",
     "0 testů pro extractor, browser, auth",
     "PRAVDA — testy pokrývají jen scoring a KBWriter. Extractor/browser/auth = 0%",
     "🔴 Nutná priorita"),
    ("F-07", "analysis/kb_writer.py (396 lines)",
     "KBWriter míchá 3 odpovědnosti (metadata, markdown, git)",
     "PRAVDA — jedna třída, 396 lines. metadata JSON + markdown report + git subprocess = 3 v jednom",
     "⏳ Plánováno (P1-4)"),
    ("F-08", "analysis/kb_writer.py:380-396 (commit_changes)",
     "Git commit blokuje event loop",
     "PRAVDA — subprocess.run() je synchronní uvnitř async handleru. Bez asyncio.to_thread()",
     "⏳ Plánováno (P1-5)"),
    ("F-09", "server.py / tools/job.py",
     "L2 Resources nejsou implementovány",
     "PRAVDA — chybí @mcp.resource dekorátory. Server má jen L1 tools",
     "⏳ Plánováno (P2-1)"),
    ("F-10", "server.py",
     "Chybí @mcp.prompt()",
     "PRAVDA — žádný prompt template není definován",
     "⏳ Plánováno (P3-1)"),
    ("F-11", "tools/job.py:210-261 (analyze_saved_jobs loop)",
     "Chybí report_progress()",
     "PRAVDA — smyčka používá ctx.info(), ne ctx.report_progress()",
     "⏳ Plánováno (P2-2)"),
    ("F-12", "analysis/config.py (FAKE_ENGINEER_KEYWORDS)",
     "Pouze 8 termínů — tenké pokrytí",
     "PRAVDA — 8 termínů je málo. role_score je binární (title_has_engineering × text_has_fake)",
     "⏳ Plánováno (P3-3)"),
    ("F-13", "scraping/extractor.py:395-399 (selector fallback)",
     "Fallback 'selector not found' bez metriky četnosti",
     "PRAVDA — logger.info() bez počítání výskytů. Nelze detekovat drift",
     "⏳ Plánováno (P3-4)"),
    ("F-14", "cli.py (_login, _status)",
     "asyncio.run() uvnitř — riziko při znovupoužití",
     "PRAVDA — asyncio.run() nelze volat z běžící event loop. Latentní riziko",
     "✅ Info — nízká priorita"),
    ("F-15", "analysis/config.py (262 lines)",
     "Keywords jako Python literály místo YAML",
     "ČÁSTEČNĚ — Python je typově bezpečnější, ale hůře editovatelný. Audit doporučuje ponechat",
     "✅ Konsensus — ponechat"),
    ("F-16", "— (chybějící LICENSE soubor)",
     "Chybí LICENSE",
     "PRAVDA — žádný LICENSE soubor v rootu repa",
     "⏳ Plánováno (P2-4)"),
    ("F-17", "pyproject.toml",
     "Balíček není publikovatelný na PyPI",
     "ČÁSTEČNĚ — pyproject.toml existuje, ale [project] sekce chybí (je tam custom formát uv)",
     "⏳ Plánováno (P2-4)"),
    ("F-18", "— (chybějící Dockerfile)",
     "Chybí Dockerfile pro stdio/SSE nasazení",
     "PRAVDA — žádný Dockerfile neexistuje",
     "⏳ Plánováno (P3-2)"),
    ("F-19", "analysis/report_generator.py",
     "Hodnoty dimenzí neodpovídají diskrétním hodnotám scoreru",
     "NEPOTVRZENO — report_generator.py:545 byl v podkladech jen jako shrnutí. Je třeba ověřit transformaci",
     "🔴 Nutná verifikace"),
]

for row_data in rows_data:
    row = table.add_row().cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[3].text = row_data[3]
    row[4].text = row_data[4]

doc.add_paragraph()

doc.add_page_break()

# ── 3. Debugging playbook ──────────────────────────────────────────────
doc.add_heading("3. Debugging Playbook — krok za krokem", level=1)

doc.add_paragraph(
    "Cílem playbooku je dovést linkedin-mcp-custom z aktuálního stavu (4/50 úspěšnost, "
    "health score 5/10) do produkčního stavu s ≥80% úspěšností pipeline a publikovatelné "
    "verze v0.2.0."
)

doc.add_heading("3.1 Fáze 0: Diagnostika (1 den)", level=2)
steps_0 = [
    "Ověřit F-02: Prostudovat run_pipeline.py:410-426 a potvrdit, že asyncio.wait_for obaluje celý task včetně čekání na semafor (nikoli až po acquire). Fix: přesunout wait_for dovnitř async with semaphore.",
    "Ověřit F-04: Porovnat prahy v scorer.py (65/50/40/0) s prahy v config_snapshot (70/45/30). Zjistit, která sada je skutečně použita. Zkontrolovat surový pipeline JSON pro job 4439059407.",
    "Ověřit F-19: Prostudovat scénu mezi scorer.py→EROIResult.to_dict()→report_generator.py. Zjistit, kde vzniká hodnota 20.0% pro dimenzi formal (scorer vrací jen 10/30/50).",
    "Spustit pipeline s max_concurrent=1 (sekvenčně) a ověřit, zda klesne počet timeoutů.",
    "Spustit pipeline s per_job_timeout_ms=30000 (progressive timeout již implementován, ověřit chování).",
    "Zalogovat tři časové značky na úlohu: task creation, semaphore acquire, goto start.",
]
for i, s in enumerate(steps_0, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"{i}. ")
    run.bold = True
    p.add_run(s)

doc.add_heading("3.2 Fáze 1: Pipeline overhaul (2-3 dny)", level=2)
steps_1 = [
    "Fix F-02: Přesunout asyncio.wait_for dovnitř async with semaphore v run_pipeline.py, aby timeout měřil pouze dobu práce, ne dobu čekání ve frontě.",
    "Vylepšit P2 backoff: Přidat rolling-window (posledních N=5 výsledků). Při >50% fail rate v okně: exponenciální pauza 30s→60s→120s s resetem po úspěchu.",
    "Upravit výchozí max_concurrent na 1 (sekvenční mód s náhodnou prodlevou 2-5s). Ponechat paralelní režim jako explicitní --fast flag.",
    "Ověřit, že progressive timeout (extractor.py) správně interaguje s pipeline-level backoffem (neduplicitní logika).",
    "Přidat metriku četnosti selector fallbacku (F-13) — při >80% fallback rate v jednom scrape run: log jako ERROR + anomaly.",
]
for i, s in enumerate(steps_1, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"{i}. ")
    run.bold = True
    p.add_run(s)

doc.add_heading("3.3 Fáze 2: EROI scoring opravy (1 den)", level=2)
steps_2 = [
    "Vyřešit F-04: Sjednotit prahy mezi config.yaml a scorer.py. Rozhodnout, která sada je autoritativní. Doplnit unit test, který ověřuje konzistenci.",
    "Vyřešit F-05: Rozšířit growth_score(company) → growth_score(company, text). Přidat keyword scan pro růstové signály (Series A/B/C, fast-growing, scale-up, rozšiřujeme tým).",
    "Vyřešit F-19: Zprůhlednit transformaci dimenzionálních skóre do procent. Doplnit dokumentaci k EROIResult.to_dict().",
    "Rozšířit FAKE_ENGINEER_KEYWORDS (F-12) z 8 na ~20 termínů na základě logovaných neshod.",
]
for i, s in enumerate(steps_2, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"{i}. ")
    run.bold = True
    p.add_run(s)

doc.add_heading("3.4 Fáze 3: Refactoring (2 dny)", level=2)
steps_3 = [
    "Refaktor KBWriter (F-07): Rozdělit na MetadataStore, MarkdownReportWriter, GitCommitter. Zachovat KBWriter jako fasádu.",
    "Async-safe git (F-08): Nahradit subprocess.run() → asyncio.create_subprocess_exec() nebo asyncio.to_thread().",
    "Odstranit dead code: _track_nav_task (browser.py:155) není nikde voláno.",
    "Fix race condition: Přidat asyncio.Lock kolem get_or_create_browser().",
    "Fix route inheritance: Aplikovat resource blocking i na pool pages z create_page().",
    "Fix get_saved_jobs timeout: Přidat asyncio.wait_for() s timeoutem do MCP tool handleru.",
    "Centralizovat VERDICT_ICON do jediného zdroje (schemas.py nebo config.py).",
]
for i, s in enumerate(steps_3, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"{i}. ")
    run.bold = True
    p.add_run(s)

doc.add_heading("3.5 Fáze 4: Testování (2 dny)", level=2)
steps_4 = [
    "Přidat unit testy pro extractor.py: AsyncMock(spec=Page) s fixture daty. Testovat _extract_job_metadata, _extract_job_ids, navigate_to_page (úspěch, timeout, checkpoint, prázdný obsah).",
    "Přidat unit testy pro browser.py: Mockovat Patchright, testovat pool round-robin, close_browser drain, create_page.",
    "Přidat unit testy pro auth.py: Mockovat Page, testovat check_cached_auth, ensure_authenticated, checkpoint detekci.",
    "Lokální E2E fixture server: aiohttp test server obsluhující sanitizované HTML snímky LinkedIn stránek. Patchright naviguje na localhost:PORT.",
    "Kombinovat s AsyncMock(spec=LinkedInExtractor) pro tools/job.py testy — odchytit signaturový drift.",
    "CI gate na testy: GitHub Actions workflow musí failovat při pádu testů (dnes jen produkuje outputs).",
]
for i, s in enumerate(steps_4, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"{i}. ")
    run.bold = True
    p.add_run(s)

doc.add_page_break()

# ── 4. Produkční readiness ────────────────────────────────────────────
doc.add_heading("4. Produkční readiness", level=1)

doc.add_heading("4.1 Požadavky pro v0.2.0 release", level=2)

doc.add_paragraph("Minimální kritéria pro release (přísná NE-GO při nesplnění):")

reqs = [
    ("P0-1", "Fix F-02: per-job timeout ukotvený za semaforem", "HIGH"),
    ("P0-2", "Pipeline-level backoff (rolling-window)", "HIGH"),
    ("P0-3", "HTTP status logging z goto() — implementováno", "DONE"),
    ("P0-4", "Verdiktová konzistence (F-04) — ověřeno a opraveno", "HIGH"),
    ("P1-1", "Výchozí max_concurrent=1 s náhodnou prodlevou", "MEDIUM"),
    ("P1-2", "Sekvenční fallback při >50% selhání — implementováno", "DONE"),
    ("P1-3", "Test coverage extractor.py + browser.py + auth.py", "HIGH"),
    ("P2-4", "LICENSE, oprava pyproject.toml, .gitignore pro profil", "HIGH"),
]

t2 = doc.add_table(rows=1, cols=3)
t2.style = "Light Grid Accent 1"
h2 = t2.rows[0].cells
h2[0].text = "ID"
h2[1].text = "Požadavek"
h2[2].text = "Priorita"

for r in reqs:
    row = t2.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
    row[2].text = r[2]

doc.add_heading("4.2 Provozní metrika úspěšnosti", level=2)
doc.add_paragraph(
    "Cíl pro v0.2.0: ≥80% úspěšnost pipeline (40/50 jobs scored) při běhu v pracovních hodinách "
    "(9:00-17:00). Měřeno na 3 po sobě jdoucích runech. Pokud je LinkedIn rate-limiting "
    "neodstranitelný (endpoint blokuje i sekvenční přístup s náhodnými prodlevami), "
    "dokumentovat očekávanou úspěšnost jako known limitation v README."
)

doc.add_heading("4.3 LinkedIn scraping strategie — finální doporučení", level=2)
recs = [
    "Výchozí mód: sekvenční (max_concurrent=1) s náhodnou prodlevou 2-5s mezi úlohami.",
    "Rychlý mód (--fast): max_concurrent=2-3, ale s varováním o zvýšeném riziku rate-limitingu.",
    "Pipeline-level backoff: rolling-window N=5, >50% fail rate → 30s→60s→120s pauza.",
    "HTTP status logging: již implementováno v _retry_goto (debug level).",
    "Persistentní profil: zachovat, cookies přežívají restart. Dokumentovat jako bezpečnostní riziko.",
]
for r in recs:
    doc.add_paragraph(r, style="List Bullet")

doc.add_page_break()

# ── 5. Publish readiness ──────────────────────────────────────────────
doc.add_heading("5. Publish readiness", level=1)

doc.add_heading("5.1 Blokující položky", level=2)
pub_blocking = [
    "Doplnit LICENSE (MIT nebo Apache-2.0) — odhad 15 min.",
    "Opravit pyproject.toml na standardní PEP 621 formát — [project], [project.scripts], [build-system]. Balíček musí být instalovatelný přes pip install linkedin-mcp-custom.",
    "Ověřit .gitignore: ~/.linkedin-mcp-custom/profile/ nikdy nesmí být v gitu. Profil obsahuje reálné LinkedIn session cookies.",
    "Doplnit SECURITY.md s upozorněním, že automatizace LinkedIn bez oficiálního API může porušovat Terms of Service.",
]
for it in pub_blocking:
    doc.add_paragraph(it, style="List Bullet")

doc.add_heading("5.2 Nice-to-have", level=2)
pub_nice = [
    "Dockerfile: multi-stage build, profil prohlížeče jako mountovaný volume.",
    "CONTRIBUTING.md + issue templates (bug_report.md, feature_request.md).",
    "CI gate: GitHub Actions musí failovat při pádu testů.",
    "README badges: tests passing, python version, license.",
    "Dokumentace očekávané úspěšnosti pipeline a známých limitací LinkedIn scrapingu.",
]
for it in pub_nice:
    doc.add_paragraph(it, style="List Bullet")

doc.add_heading("5.3 Bezpečnostní checklist", level=2)
sec_items = [
    "Profil prohlížeče (~/.linkedin-mcp-custom/profile/) je v .gitignore? Ověřit.",
    "Session cookies nejsou logovány ani expozovány v MCP tool výstupech? Ověřit.",
    "KBWriter zapisuje do B2B-Knowledge-Base — riziko unintentional commit citlivých dat? Dokumentovat.",
    "CLI --login mód spouští headless=false — vhodné pro první setup. Dokumentovat.",
    "Zvážit volitelné šifrování profilu at-rest jako obranu do hloubky.",
]
for it in sec_items:
    doc.add_paragraph(it, style="List Bullet")

doc.add_page_break()

# ── 6. Prioritizovaná road mapa ──────────────────────────────────────
doc.add_heading("6. Prioritizovaná road mapa k v0.2.0", level=1)

doc.add_paragraph("Odvozeno z auditu (kap. 12) + vlastní korelace. Časové odhady jsou kumulativní.")

roadmap = [
    ("Týden 1", "Diagnostika + Fáze 0",
     "Ověřit F-02, F-04, F-19. Spustit pipeline s max_concurrent=1. Zalogovat časové značky. "
     "Potvrdit nebo vyvrátit hypotézy auditu."),
    ("Týden 1-2", "Fáze 1: Pipeline overhaul",
     "Fix F-02 (per-job timeout). Vylepšit P2 backoff (rolling-window). "
     "Zprovoznit sekvenční mód jako výchozí. Ověřit progressive timeout chování."),
    ("Týden 2", "Fáze 2: EROI scoring opravy",
     "Fix F-04 (verdict konzistence). Fix F-05 (growth_score + text). "
     "Fix F-19 (transparence dimenzí). Rozšířit FAKE_ENGINEER_KEYWORDS."),
    ("Týden 2-3", "Fáze 3: Refactoring",
     "Refaktor KBWriter. Async-safe git. Odstranit dead code. "
     "Fix race condition v browser.py. Fix route inheritance."),
    ("Týden 3-4", "Fáze 4: Testování",
     "AsyncMock testy pro extractor/browser/auth. Lokální E2E fixture server. CI gate."),
    ("Týden 4", "Publish readiness",
     "LICENSE, pyproject.toml, SECURITY.md, .gitignore. Dockerfile (nice-to-have). "
     "Release v0.2.0."),
]

t3 = doc.add_table(rows=1, cols=3)
t3.style = "Light Grid Accent 1"
h3 = t3.rows[0].cells
h3[0].text = "Časový rámec"
h3[1].text = "Fáze"
h3[2].text = "Náplň"

for r in roadmap:
    row = t3.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
    row[2].text = r[2]

doc.add_paragraph()

doc.add_heading("6.1 Go/No-Go kritéria pro v0.2.0 release", level=2)
gongo = [
    "GO: ≥80% pipeline success rate (40/50) ve 3 po sobě jdoucích runech v pracovní době.",
    "GO: Všechny P0 položky vyřešeny + testovány.",
    "NO-GO: Přetrvávající verdiktová nesrovnalost (F-04) bez jasného vysvětlení.",
    "NO-GO: Pipeline stále padá na >50% při sekvenčním módu s prodlevami.",
    "NO-GO: Chybí LICENSE, SECURITY.md, nebo není ošetřen profil prohlížeče v .gitignore.",
]
for g in gongo:
    doc.add_paragraph(g, style="List Bullet")

# ── Save ───────────────────────────────────────────────────────────────
output_path = Path(__file__).resolve().parent.parent / "docs" / "audit_analyza_a_playbook_v0.2.0.docx"
doc.save(str(output_path))
print(f"OK: {output_path}")
